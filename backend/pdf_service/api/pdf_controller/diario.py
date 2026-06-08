from reportlab.lib import colors
from reportlab.lib.colors import Color
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Spacer,
    Paragraph,
    Flowable,
    Table,
    TableStyle
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas
from reportlab.graphics.shapes import Rect, Line, Drawing
from reportlab.lib.utils import ImageReader
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from datetime import datetime
import docx
import unicodedata
from collections import defaultdict
from io import BytesIO
from api.config import FONTS_DIR,DATA_DIR
import os
import pdfplumber
from io import BytesIO
import re
import fitz  # PyMuPDF
from PIL import Image as PILImage
import numpy as np
from urllib.parse import urlparse
from xml.sax.saxutils import escape

def create_header_template(context):
    def header_template(canvas: canvas.Canvas, doc):
        width, height = A4
        canvas.saveState()

        # Caminho da imagem do brasão
        brasao_bytes = context.get('brasao')
        with BytesIO(brasao_bytes) as brasao:
            try:    
                image = ImageReader(brasao)
                image_width = 3 * cm
                image_height = 3.5 * cm

                # Posição do brasão à esquerda
                left_image_x = 1 * cm
                left_image_y = height - 4 * cm  # Alinha com o título

                # Posição do brasão à direita
                right_image_x = width - 4 * cm
                right_image_y = left_image_y

                # Desenha o brasão à esquerda
                canvas.drawImage(image,left_image_x,left_image_y,width=image_width,height=image_height,preserveAspectRatio=True, mask='auto')

                # Desenha o brasão à direita
                canvas.drawImage(image, right_image_x, right_image_y, width=image_width, height=image_height, preserveAspectRatio=True, mask='auto')
            except Exception as e:
                print(f"Erro ao carregar a imagem: {e}")

        # Ajuste para o título principal com fonte maior e negrito
        canvas.setFont("Impact", 30)
        canvas.setFillColor("#03A64A")
        canvas.drawCentredString(width / 2.0, height - 1.8 * cm, "DIÁRIO OFICIAL ELETRÔNICO")

        # Subtítulo
        canvas.setFont("Impact", 16)
        canvas.drawCentredString(width / 2.0, height - 2.7 * cm, f"{str(context.get('tipo','')).upper()} DE {str(context.get('cidade','')).upper()}")

        # Texto adicional
        canvas.setFont("Calibri-Bold", 8)
        canvas.setFillColor(colors.black)
        canvas.drawCentredString(width / 2.0, height - 3.3 * cm, f"LEI MUNICIPAL N° {str(context.get('num_lei','')).upper()} DE {context.get('data_lei','')}",charSpace=3,wordSpace=5)

        canvas.setStrokeColor(colors.gray)
        canvas.setLineWidth(0.7)
        canvas.line(4.4 * cm, height - 3.5 * cm, width - 4.4 * cm, height - 3.5 * cm)

        # Informações de data
        canvas.setFont("Calibri-Bold", 9)
        canvas.drawCentredString(width / 2.0, height - 3.8 * cm, f"ANO {str(context.get('ano_romano','')).upper()} - {str(context.get('cidade','')).upper()}, {str(context.get('data-publicacao','')).upper()} - Nº {context.get('edicao','')}")

        # Linha horizontal no final do cabeçalho
        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(2)
        canvas.line(1 * cm, height - 4.2 * cm, width - 1 * cm, height - 4.2 * cm)

        canvas.restoreState()

    return header_template

def create_container_responsavel(context):
    def container_responsavel(canvas: canvas.Canvas, doc):
        page_width, page_height = A4
        margin = 1 * cm

        # Dimensões do container principal
        rect_x = margin
        rect_y = margin + 0.5 * cm
        rect_width = (page_width / 2) - margin - 0.5 * cm
        rect_height = 4.5 * cm

        canvas.setStrokeColor(colors.gray)
        canvas.setLineWidth(1)
        canvas.roundRect(rect_x, rect_y, rect_width, rect_height, radius=10, stroke=1, fill=0)

        inner_rect_margin = 0.2 * cm
        inner_rect_x = rect_x + inner_rect_margin
        inner_rect_y = rect_y + inner_rect_margin
        inner_rect_width = rect_width - 2 * inner_rect_margin
        inner_rect_height = rect_height - 2 * inner_rect_margin
        brasao_bytes = context.get('brasao')
        with BytesIO(brasao_bytes) as brasao:
            try:
                image = ImageReader(brasao)
                image_width = 2 * cm
                image_height = 2 * cm
                image_x = inner_rect_x + (inner_rect_width - image_width) / 2
                image_y = inner_rect_y + inner_rect_height - image_height
                canvas.drawImage(image, image_x, image_y, width=image_width, height=image_height, preserveAspectRatio=True, mask='auto')
            except Exception as e:
                print(f"Erro ao carregar a imagem: {e}")
        grey_rect_x = inner_rect_x
        grey_rect_y = inner_rect_y
        grey_rect_width = inner_rect_width
        grey_rect_height = 2 * cm

        canvas.setFillColor(colors.gray)
        canvas.roundRect(grey_rect_x, grey_rect_y, grey_rect_width, grey_rect_height, radius=5, stroke=0, fill=1)

        # Texto dentro do container
        text_y = grey_rect_y + 0.5 * cm
        canvas.setFont("Calibri-Bold", 10)
        canvas.setFillColor(colors.black)
        canvas.drawCentredString(grey_rect_x + (grey_rect_width / 2), text_y + 0.6 * cm, context.get('responsavel',''))

        canvas.setFont("Calibri", 8)
        canvas.drawCentredString(grey_rect_x + (grey_rect_width / 2), text_y, context.get('cargo',''))

        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(1)
        canvas.line(1 * cm, margin, page_width - 1 * cm, margin)
    return container_responsavel

def create_header_template_later(context):
    def header_template_later(canvas:canvas.Canvas, doc):
        width, height = A4
        canvas.saveState()
        canvas.setFont("Calibri-Bold", 8)
        canvas.drawString(1.5 * cm, height - 0.8 * cm, f"DÍARIO OFICIAL {str(context.get('tipo','')).upper()} DE {str(context.get('cidade','')).upper()}")
        canvas.drawString(width / 2.0 + 1 * cm, height - 0.8 * cm, f"ANO {str(context.get('ano_romano','')).upper()} - Nº {context.get('edicao','')} - {str(context.get('data-publicacao','')).upper()}")
        canvas.setStrokeColor(colors.black)
        canvas.setLineWidth(1)
        canvas.line(1 * cm, height - 1 * cm, width - 1 * cm, height - 1 * cm)

        canvas.restoreState()
    return header_template_later

def footer_template(canvas:canvas.Canvas, doc):
    page_width, page_height = A4
    margin = 1 * cm
    canvas.setStrokeColor(colors.black)
    canvas.setLineWidth(1)
    canvas.line(1 * cm, margin, page_width - 1 * cm, margin)

class SectionTitle(Flowable):
    def __init__(self, text, width, height, min_height=None, offset_x=0):
        super().__init__()
        self.text = text
        self.width = width
        self.base_height = height
        self.min_height = min_height or height
        self.offset_x = offset_x  # Deslocamento horizontal para a esquerda
        
        # Estilo do parágrafo para medir e renderizar o texto
        self.text_style = ParagraphStyle(
            'SectionTitleText',
            fontName="Calibri-Bold",
            fontSize=12,
            leading=14,
            alignment=TA_CENTER,
            textColor=colors.white
        )
        
        # Calcula a altura necessária baseada no texto
        self._calculate_height()
    
    def _calculate_height(self):
        """Calcula a altura necessária para acomodar o texto com quebra de linha"""
        para = Paragraph(self.text, self.text_style)
        # Calcula o espaço necessário (largura disponível com margem)
        available_width = self.width - 0.4 * cm  # margem lateral
        w, h = para.wrap(available_width, 1000)  # altura grande para não limitar
        
        # Altura = altura do texto + padding vertical
        padding = 0.4 * cm
        calculated_height = h + padding
        
        # Usa o maior entre a altura calculada e a altura base/mínima
        self.height = max(calculated_height, self.min_height)

    def wrap(self, aW, aH):
        """Retorna as dimensões do flowable"""
        return (self.width, self.height)

    def draw(self):
        # Configuração do fundo azul e linha cinza
        blue_bg = Color(0, 0.5, 1)  # Azul
        gray_line = Color(0.6, 0.6, 0.6)  # Cinza claro

        # Desenho do retângulo azul como fundo (com offset_x)
        self.canv.setFillColor(blue_bg)
        self.canv.rect(self.offset_x, 0, self.width, self.height, stroke=0, fill=1)

        # Cria o parágrafo com quebra de linha automática
        para = Paragraph(self.text, self.text_style)
        available_width = self.width - 0.4 * cm
        w, h = para.wrap(available_width, self.height)
        
        # Centraliza verticalmente
        y_offset = (self.height - h) / 2
        
        # Desenha o parágrafo (com offset_x)
        para.drawOn(self.canv, self.offset_x + 0.2 * cm, y_offset)

        # Linha cinza abaixo do retângulo azul (com offset_x)
        self.canv.setStrokeColor(gray_line)
        self.canv.setLineWidth(1)
        self.canv.line(self.offset_x, -0.1 * cm, self.offset_x + self.width, -0.1 * cm)

class DocumentImage(Flowable):
    """
    Flowable personalizado para exibir imagens extraídas dos documentos.
    Mantém proporção e se ajusta automaticamente à largura da coluna.
    """
    def __init__(self, image_data, max_width, max_height=None, alignment='center'):
        super().__init__()
        self.image_data = image_data
        self.max_width = max_width
        self.max_height = max_height or max_width  # Se não especificado, usa largura como altura máxima
        self.alignment = alignment
        
        # Calcula as dimensões da imagem
        self._calculate_dimensions()
    
    def _calculate_dimensions(self):
        """Calcula as dimensões finais da imagem mantendo a proporção"""
        try:
            # Usa PIL para obter dimensões originais da imagem
            from PIL import Image as PILImage
            image = PILImage.open(BytesIO(self.image_data))
            original_width, original_height = image.size
            
            # Calcula fator de escala para ajustar à largura máxima
            scale_factor_width = self.max_width / original_width
            scale_factor_height = self.max_height / original_height
            
            # Usa o menor fator para manter a proporção
            scale_factor = min(scale_factor_width, scale_factor_height)
            
            # Aplica um fator de aumento de 120% para imagens maiores
            scale_factor = scale_factor * 1.25
            
            self.width = original_width * scale_factor
            self.height = original_height * scale_factor
            
        except Exception as e:
            print(f"Erro ao calcular dimensões da imagem: {e}")
            # Valores padrão em caso de erro
            self.width = min(self.max_width, 5 * cm)
            self.height = min(self.max_height, 5 * cm)
    
    def draw(self):
        """Desenha a imagem no canvas"""
        try:
            # Cria ImageReader a partir dos dados da imagem
            image_reader = ImageReader(BytesIO(self.image_data))
            
            # Calcula posição X baseada no alinhamento
            if self.alignment == 'center':
                x = (self.max_width - self.width) / 2
            elif self.alignment == 'right':
                x = self.max_width - self.width
            else:  # left
                x = 0
            
            # Desenha a imagem
            self.canv.drawImage(
                image_reader, 
                x, 0, 
                width=self.width, 
                height=self.height,
                preserveAspectRatio=True, 
                mask='auto'
            )
            
        except Exception as e:
            print(f"Erro ao desenhar imagem: {e}")
            # Em caso de erro, desenha um retângulo placeholder
            self.canv.setStrokeColor(colors.red)
            self.canv.setFillColor(colors.lightgrey)
            self.canv.rect(0, 0, self.width, self.height, stroke=1, fill=1)
            
            # Adiciona texto de erro
            self.canv.setFont("Calibri", 8)
            self.canv.setFillColor(colors.red)
            self.canv.drawCentredString(self.width / 2, self.height / 2, "Erro ao carregar imagem")


class SplittableDocumentImage(Flowable):
    """
    Flowable que pode dividir uma imagem se ela não couber na página/coluna atual.
    Retorna a parte que não coube para ser inserida na próxima página.
    """
    def __init__(self, image_data, max_width, max_height=None, alignment='center', scale_factor=1.25):
        super().__init__()
        self.image_data = image_data
        self.image_data_original = image_data  # Bytes originais desta imagem
        self.max_width = max_width
        self.max_height = max_height or max_width
        self.alignment = alignment
        self.scale_factor = scale_factor
        # Calcula dimensões
        self._calculate_dimensions()
    
    def _calculate_dimensions(self):
        """Calcula as dimensões finais da imagem mantendo a proporção"""
        try:
            image = PILImage.open(BytesIO(self.image_data))
            original_width, original_height = image.size
            
            # Escala baseada na largura (preserva proporção)
            scale_w = self.max_width / original_width
            scale_h = (self.max_height / original_height) if self.max_height else scale_w
            scale = min(scale_w, scale_h) * self.scale_factor

            self.width = original_width * scale
            self.height = original_height * scale

            # Guarda dimensões e escala correntes
            self.original_width = original_width
            self.original_height = original_height
            self.current_scale = scale
            
        except Exception as e:
            print(f"Erro ao calcular dimensões da imagem: {e}")
            self.width = min(self.max_width, 5 * cm)
            self.height = min(self.max_height, 5 * cm)
            self.original_width = self.width
            self.original_height = self.height
            self.current_scale = 1.0
    
    def wrap(self, aW, aH):
        # Nunca muta estado aqui. Apenas informa o tamanho do flowable.
        # Reserva a largura da coluna para alinhamento adequado.
        return (self.max_width, self.height)
    
    def split(self, aW, aH):
        """Divide a imagem quando não cabe; retorna [top, bottom] sem mutar self."""
        # Se já cabe, não divide
        if self.height <= aH:
            return []

        # Espaço muito pequeno → não divide
        if aH < 20:
            return []

        # Converte a altura disponível (em pontos) para pixels na imagem original
        pixels_that_fit = int(max(1, (aH / self.current_scale)))

        # Busca ponto de corte preferindo linhas abaixo do alvo (preencher ao máximo)
        dynamic_range = max(60, int(self.original_height * 0.05))
        split_point = find_best_split_point(
            self.image_data_original,
            pixels_that_fit,
            search_range=dynamic_range
        )

        # Garante que o topo não ultrapasse o espaço disponível
        if split_point > pixels_that_fit:
            split_point = pixels_that_fit

        # Valida ponto de corte
        if split_point < 5 or split_point >= self.original_height - 5:
            return []

        # Corta a imagem em duas partes
        top_bytes, bottom_bytes = split_image_at_height(self.image_data_original, split_point)

        # Cria flowables independentes (não mutar self)
        top_flowable = SplittableDocumentImage(
            top_bytes,
            self.max_width,
            self.max_height,
            self.alignment,
            self.scale_factor
        )
        bottom_flowable = SplittableDocumentImage(
            bottom_bytes,
            self.max_width,
            self.max_height,
            self.alignment,
            self.scale_factor
        )

        # Segurança: o primeiro deve caber
        if top_flowable.height > aH:
            return []

        return [top_flowable, bottom_flowable]
    
    def draw(self):
        """Desenha a imagem no canvas"""
        try:
            image_reader = ImageReader(BytesIO(self.image_data))
            
            # Calcula posição X baseada no alinhamento
            if self.alignment == 'center':
                x = (self.max_width - self.width) / 2
            elif self.alignment == 'right':
                x = self.max_width - self.width
            else:
                x = 0
            
            self.canv.drawImage(
                image_reader,
                x, 0,
                width=self.width,
                height=self.height,
                preserveAspectRatio=True,
                mask='auto'
            )
            
        except Exception as e:
            print(f"Erro ao desenhar imagem: {e}")
            self.canv.setStrokeColor(colors.red)
            self.canv.setFillColor(colors.lightgrey)
            self.canv.rect(0, 0, self.width, self.height, stroke=1, fill=1)
            self.canv.setFont("Calibri", 8)
            self.canv.setFillColor(colors.red)
            self.canv.drawCentredString(self.width / 2, self.height / 2, "Erro ao carregar imagem")


def create_splittable_image(image_data, max_width, max_height=None, alignment='center', scale_factor=1.25):
    """
    Cria uma imagem que pode ser dividida automaticamente se não couber na página.
    
    Args:
        image_data: Bytes da imagem
        max_width: Largura máxima permitida
        max_height: Altura máxima permitida (opcional)
        alignment: Alinhamento da imagem ('left', 'center', 'right')
        scale_factor: Fator de escala adicional (padrão 1.25 = 125%)
    
    Returns:
        Instância de SplittableDocumentImage
    
    Exemplo de uso:
        # Ao invés de usar DocumentImage:
        # image = DocumentImage(img_bytes, column_width)
        
        # Use SplittableDocumentImage:
        image = create_splittable_image(img_bytes, column_width)
        story.append(image)
        # A imagem será automaticamente dividida se não couber!
    """
    return SplittableDocumentImage(image_data, max_width, max_height, alignment, scale_factor)


def split_paragraphs(text: str) -> list[str]:
    """
    Agrupa linhas em parágrafos usando:
      1) fim em pontuação forte (.!?;:)
      2) nova linha que COMEÇA com aspas (“ ou ")
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    paras: list[str] = []
    if not lines:
        return paras

    buf = lines[0]
    for line in lines[1:]:
        # Se a linha começa com aspas, força nova quebra
        if line.startswith('“') or line.startswith('"'):
            paras.append(buf)
            buf = line
            continue

        # Se o buffer atual termina em pontuação forte, quebra
        last = buf.rstrip()[-1]
        if last in ".!?;:":
            paras.append(buf)
            buf = line
        else:
            buf += " " + line

    paras.append(buf)
    return paras


def _consolidate_table_rows(table_data):
    """
    Consolida linhas fragmentadas em uma única linha lógica.
    Trata especificamente quebras de linha em cabeçalhos e células fragmentadas.
    """
    if not table_data or len(table_data) <= 1:
        return table_data
    
    def is_empty_cell(cell):
        """Verifica se uma célula está efetivamente vazia"""
        if cell is None:
            return True
        cell_str = str(cell).strip()
        return not cell_str or cell_str.lower() in ['none', 'null', 'nan', '']
    
    def is_sparse_row(row):
        """Verifica se uma linha é esparsa (poucos dados)"""
        non_empty_count = sum(1 for cell in row if not is_empty_cell(cell))
        return non_empty_count <= 1  # Alterado para máximo 1 célula preenchida
    
    def has_continuation_content(row):
        """Verifica se a linha tem conteúdo que parece ser continuação de cabeçalho"""
        continuation_words = ['por', 'questão', 'de', 'da', 'do', 'em', 'na', 'no', 'e', 'valor']
        for cell in row:
            if not is_empty_cell(cell):
                cell_text = str(cell).strip()
                # Verifica se é uma palavra isolada que pode ser continuação
                words = cell_text.split()
                if len(words) <= 2 and any(word in continuation_words for word in words):
                    return True
        return False
    
    def is_complete_content_row(row):
        """Verifica se uma linha tem conteúdo completo (não é continuação)"""
        for cell in row:
            if not is_empty_cell(cell):
                cell_text = str(cell).strip()
                # Se tem mais de 10 caracteres ou termina com ponto, provavelmente é conteúdo completo
                if len(cell_text) > 10 or cell_text.endswith('.') or '/' in cell_text:
                    return True
        return False
    
    consolidated = []
    i = 0
    
    while i < len(table_data):
        current_row = table_data[i][:]
        
        # Se é a primeira linha (cabeçalho), verifica se próximas linhas são continuações
        if i == 0:
            j = i + 1
            # Verifica próximas linhas esparsas que podem ser parte do cabeçalho
            while j < len(table_data) and j < i + 3:  # Máximo 2 linhas à frente
                next_row = table_data[j]
                
                # Só consolida se for realmente uma continuação de cabeçalho
                if (is_sparse_row(next_row) and 
                    has_continuation_content(next_row) and 
                    not is_complete_content_row(next_row)):
                    
                    # Consolida com o cabeçalho
                    for k, (curr_cell, next_cell) in enumerate(zip(current_row, next_row)):
                        if not is_empty_cell(next_cell):
                            if is_empty_cell(curr_cell):
                                current_row[k] = next_cell
                            else:
                                # Ambos têm conteúdo, concatena
                                curr_text = str(curr_cell).strip()
                                next_text = str(next_cell).strip()
                                current_row[k] = f"{curr_text}\n{next_text}"
                    j += 1
                else:
                    break
            
            consolidated.append(current_row)
            i = j
            continue
        
        # Para linhas de dados, apenas consolida se for realmente uma continuação fragmentada
        j = i + 1
        if j < len(table_data):
            next_row = table_data[j]
            
            # Conta células não vazias
            next_non_empty = sum(1 for cell in next_row if not is_empty_cell(cell))
            current_non_empty = sum(1 for cell in current_row if not is_empty_cell(cell))
            
            # Verifica se a próxima linha parece ser uma continuação de texto fragmentado
            is_text_continuation = False
            if next_non_empty == 1:  # Próxima linha tem apenas uma célula
                for k, cell in enumerate(next_row):
                    if not is_empty_cell(cell):
                        cell_text = str(cell).strip()
                        # Se é um texto curto que parece ser continuação (como "1990.")
                        # e a linha atual tem conteúdo na mesma coluna ou próxima
                        if len(cell_text) <= 10:  # Removida a condição is_complete_content_row
                            # Verifica se há conteúdo relacionado na linha atual
                            for curr_k in range(max(0, k-1), min(len(current_row), k+2)):
                                if not is_empty_cell(current_row[curr_k]):
                                    curr_text = str(current_row[curr_k]).strip()
                                    # Se o texto atual termina de forma que parece incompleto
                                    if (len(curr_text) > 10 and 
                                        (curr_text.endswith('de') or 
                                         curr_text.endswith('n°') or
                                         not curr_text.endswith('.') and 
                                         any(word in curr_text.lower() for word in ['lei', 'federal', 'julho', 'decreto']))):
                                        # E se a próxima célula parece ser continuação (ano, ponto, etc)
                                        if (cell_text.endswith('.') or 
                                            cell_text.isdigit() or
                                            len(cell_text) <= 6):
                                            is_text_continuation = True
                                            break
                            if is_text_continuation:
                                break
            
            # Só consolida se for realmente uma continuação de texto
            if is_text_continuation:
                # Combina as linhas de dados
                for k in range(len(current_row)):
                    if k < len(next_row) and not is_empty_cell(next_row[k]):
                        next_content = str(next_row[k]).strip()
                        if not is_empty_cell(current_row[k]):
                            current_content = str(current_row[k]).strip()
                            # Adiciona espaço entre as partes se necessário
                            if not current_content.endswith('-'):
                                current_row[k] = f"{current_content} {next_content}"
                            else:
                                current_row[k] = f"{current_content}{next_content}"
                        else:
                            current_row[k] = next_row[k]
                i = j + 1
            else:
                i += 1
        else:
            i += 1
        
        consolidated.append(current_row)
    
    return consolidated


def _consolidate_fragmented_headers(table_data):
    """
    Consolida cabeçalhos que foram fragmentados em múltiplas linhas.
    Corrige o vazamento de conteúdo onde texto de cabeçalho aparece em linhas separadas.
    """
    if not table_data or len(table_data) < 2:
        return table_data
    
    corrected_table = []
    
    # Primeira linha (cabeçalhos) - precisa ser consolidada
    header_row = table_data[0][:]  # Cópia da primeira linha
    
    # Procura por conteúdo fragmentado nas linhas seguintes que deveria estar no cabeçalho
    fragment_content = []
    lines_to_remove = []
    
    for i in range(1, len(table_data)):
        row = table_data[i]
        
        # Verifica se esta linha contém apenas fragmentos de cabeçalho
        non_empty_cells = []
        fragment_positions = []
        
        for j, cell in enumerate(row):
            if cell and str(cell).strip():
                content = str(cell).strip()
                non_empty_cells.append((j, content))
                
                # Identifica fragmentos típicos de cabeçalho (palavras comuns em formulários)
                fragment_keywords = [
                    'serviço', 'endereço', 'telefone', 'nome', 'completo', 'chefe', 'imediato',
                    'data', 'início', 'término', 'período', 'atividades', 'desenvolvidas'
                ]
                
                if any(word in content.lower() for word in fragment_keywords):
                    fragment_positions.append((j, content))
        
        # Se a linha tem poucos elementos e são fragmentos, ela deve ser consolidada
        if len(non_empty_cells) <= 2 and fragment_positions and len(non_empty_cells) == len(fragment_positions):
            fragment_content.extend(fragment_positions)
            lines_to_remove.append(i)
        else:
            # Linha com dados reais ou vazia (formulário) - parar de procurar fragmentos
            break
    
    # Consolida fragmentos no cabeçalho
    if fragment_content:
        # Agrupa fragmentos por posição de coluna
        fragments_by_position = {}
        for pos, content in fragment_content:
            if pos not in fragments_by_position:
                fragments_by_position[pos] = []
            fragments_by_position[pos].append(content)
        
        # Adiciona fragmentos aos cabeçalhos correspondentes
        for pos, fragments in fragments_by_position.items():
            if pos < len(header_row):
                if header_row[pos] and str(header_row[pos]).strip():
                    # Cabeçalho já existe, adiciona os fragmentos
                    header_row[pos] = str(header_row[pos]).strip() + "\n" + " ".join(fragments)
                else:
                    # Posição vazia, adiciona os fragmentos
                    header_row[pos] = " ".join(fragments)
    
    corrected_table.append(header_row)
    
    # Adiciona apenas as linhas que não são fragmentos
    for i in range(1, len(table_data)):
        if i not in lines_to_remove:
            corrected_table.append(table_data[i])
    
    return corrected_table


def _remove_merged_header_rows(table_data):
    """
    Remove linhas que contêm títulos mesclados (primeira célula com conteúdo, demais vazias/None).
    Estas são linhas que deveriam ser cabeçalhos separados, não parte da tabela de dados.
    """
    if not table_data or len(table_data) <= 1:
        return table_data
    
    def is_empty_cell(cell):
        """Verifica se uma célula está vazia"""
        if cell is None:
            return True
        cell_str = str(cell).strip()
        return not cell_str or cell_str.lower() in ['none', 'null', 'nan', '']
    
    def is_merged_header_row(row):
        """
        Detecta se uma linha é um título mesclado:
        - Primeira célula tem conteúdo longo (> 10 caracteres)
        - Todas as outras células estão vazias/None
        - O conteúdo parece ser um título (contém palavras como PROCESSO, CRONOGRAMA, etc.)
        """
        if not row or len(row) < 2:
            return False
            
        # Verifica primeira célula
        first_cell = row[0]
        if is_empty_cell(first_cell):
            return False
            
        first_content = str(first_cell).strip()
        
        # Conteúdo muito longo sugere título
        if len(first_content) < 10:
            return False
        
        # Verifica se todas as outras células estão vazias
        for i in range(1, len(row)):
            if not is_empty_cell(row[i]):
                return False
        
        # Palavras-chave que indicam títulos/cabeçalhos
        title_keywords = [
            'PROCESSO', 'CRONOGRAMA', 'DADOS PESSOAIS', 
            'ESCOLARIDADE', 'CAMPANHA', 'AUTENTICAÇÃO',
            'ELEIÇÃO', 'CONSELHO', 'TUTELAR', 'SUPLEMENTAR'
        ]
        
        upper_content = first_content.upper()
        has_title_keyword = any(keyword in upper_content for keyword in title_keywords)
        
        return has_title_keyword
    
    # Remove linhas de título mesclado
    cleaned_table = []
    for row in table_data:
        if not is_merged_header_row(row):
            cleaned_table.append(row)
    
    # Se removemos todas as linhas, volta a tabela original
    if not cleaned_table:
        return table_data
    
    return cleaned_table


def _remove_duplicate_and_empty_columns(table_data):
    """
    Remove colunas duplicadas e vazias, alinhando dados com seus cabeçalhos corretos.
    """
    if not table_data or not table_data[0]:
        return table_data
    
    def is_empty_cell(cell):
        """Verifica se uma célula está vazia"""
        if cell is None:
            return True
        cell_str = str(cell).strip()
        return not cell_str or cell_str.lower() in ['none', 'null', 'nan', '']
    
    # Se há apenas uma linha (só cabeçalho), não há muito o que fazer
    if len(table_data) <= 1:
        return table_data
    
    num_cols = len(table_data[0])
    
    # 1. Identifica colunas que têm dados em pelo menos uma linha (não necessariamente o cabeçalho)
    columns_with_data = []
    for col_idx in range(num_cols):
        has_data = False
        for row_idx, row in enumerate(table_data):
            if col_idx < len(row) and not is_empty_cell(row[col_idx]):
                has_data = True
                break
        if has_data:
            columns_with_data.append(col_idx)
    
    # Se não há colunas com dados, retorna a tabela original
    if not columns_with_data:
        return table_data
    
    # 2. Para tabelas simples (2 colunas), trata de forma especial
    if num_cols == 2 and len(columns_with_data) == 2:
        # Mantém ambas as colunas mesmo se a primeira linha não tem dados na segunda coluna
        cleaned_table = []
        for row in table_data:
            new_row = []
            for col_idx in columns_with_data:
                if col_idx < len(row):
                    cell_value = row[col_idx] if not is_empty_cell(row[col_idx]) else ""
                    new_row.append(cell_value)
                else:
                    new_row.append("")
            cleaned_table.append(new_row)
        return cleaned_table
    
    # 3. Para tabelas mais complexas, usa a lógica original
    # Identifica colunas com cabeçalhos válidos
    header_columns = []
    for col_idx in range(num_cols):
        if col_idx < len(table_data[0]):
            header = table_data[0][col_idx]
            if not is_empty_cell(header):
                header_columns.append((col_idx, str(header).strip()))
    
    if not header_columns:
        # Se não há cabeçalhos válidos, remove apenas colunas completamente vazias
        cleaned_table = []
        for row in table_data:
            new_row = []
            for col_idx in columns_with_data:
                if col_idx < len(row):
                    new_row.append(row[col_idx] if not is_empty_cell(row[col_idx]) else "")
                else:
                    new_row.append("")
            cleaned_table.append(new_row)
        return cleaned_table
    
    # 4. Para cada cabeçalho, encontra os dados associados nas linhas seguintes
    aligned_columns = []
    
    for header_col, header_text in header_columns:
        column_data = [header_text]  # Começa com o cabeçalho
        
        # Para cada linha de dados
        for row_idx in range(1, len(table_data)):
            row = table_data[row_idx]
            found_data = ""
            
            # Tratamento especial para linhas que contêm "TOTAL"
            has_total = any(cell and "TOTAL" in str(cell).upper() for cell in row if not is_empty_cell(cell))
            
            if has_total:
                # Para linha de TOTAL, usa lógica especial
                if header_col == header_columns[0][0]:  # Primeira coluna
                    found_data = "TOTAL"
                elif header_col == header_columns[-1][0]:  # Última coluna (compara com a coluna original)
                    # Procura o maior valor numérico na linha (provavelmente o total)
                    max_value = ""
                    max_numeric = 0
                    for cell in row:
                        if cell and str(cell).strip() and str(cell).upper() != "TOTAL":
                            cell_value = str(cell).strip()
                            try:
                                numeric_val = float(cell_value.replace(',', '.'))
                                if numeric_val > max_numeric:
                                    max_numeric = numeric_val
                                    max_value = cell_value
                            except:
                                pass
                    found_data = max_value
                else:
                    # Para colunas do meio em linhas de TOTAL, deixa vazio para simular mesclagem
                    found_data = ""
            else:
                # Tratamento normal para linhas que não são TOTAL
                # Procura dados próximos ao cabeçalho (tolerância de algumas colunas)
                search_range = range(max(0, header_col - 2), min(len(row), header_col + 3))
                
                for search_col in search_range:
                    if search_col < len(row) and not is_empty_cell(row[search_col]):
                        candidate_data = str(row[search_col]).strip()
                        
                        # Verifica se este dado ainda não foi usado por outro cabeçalho
                        already_used = False
                        for existing_col in aligned_columns:
                            if row_idx < len(existing_col) and existing_col[row_idx] == candidate_data:
                                already_used = True
                                break
                        
                        # Para linhas TOTAL, não permite que valores numéricos sejam pegos por colunas que não são a última
                        if has_total and candidate_data.replace(',', '.').replace('.', '').isdigit():
                            if header_col != len(header_columns) - 1:  # Se não é a última coluna
                                continue  # Pula este valor
                        
                        if not already_used:
                            found_data = candidate_data
                            break
            
            column_data.append(found_data)
        
        aligned_columns.append(column_data)
    
    # 5. Verifica se há colunas com cabeçalhos duplicados e consolida
    header_to_columns = {}
    for i, column in enumerate(aligned_columns):
        header = column[0]
        if header not in header_to_columns:
            header_to_columns[header] = []
        header_to_columns[header].append((i, column))
    
    final_columns = []
    processed_indices = set()
    
    for header, col_group in header_to_columns.items():
        if len(col_group) == 1:
            # Cabeçalho único
            _, column = col_group[0]
            final_columns.append(column)
        else:
            # Cabeçalho duplicado - consolida
            consolidated_column = [header]
            max_rows = max(len(col[1]) for _, col in col_group)
            
            for row_idx in range(1, max_rows):
                row_data = []
                for _, column in col_group:
                    if row_idx < len(column) and not is_empty_cell(column[row_idx]):
                        data = str(column[row_idx]).strip()
                        if data and data not in row_data:
                            row_data.append(data)
                
                consolidated_cell = " ".join(row_data) if row_data else ""
                consolidated_column.append(consolidated_cell)
            
            final_columns.append(consolidated_column)
    
    # 6. Reconstrói a tabela transposta de volta
    if not final_columns:
        return table_data
    
    max_rows = max(len(col) for col in final_columns)
    cleaned_table = []
    
    for row_idx in range(max_rows):
        row = []
        for column in final_columns:
            if row_idx < len(column):
                cell = column[row_idx]
                row.append(cell if not is_empty_cell(cell) else "")
            else:
                row.append("")
        cleaned_table.append(row)
    
    return cleaned_table


def find_best_split_point(image_bytes: bytes, target_height: int, search_range: int = 50) -> int:
    """
    Encontra o melhor ponto para dividir uma imagem, preferencialmente em espaços em branco.
    
    Args:
        image_bytes: Bytes da imagem
        target_height: Altura alvo onde queremos dividir
        search_range: Quantos pixels acima/abaixo procurar por espaço em branco
    
    Returns:
        Altura ideal para o corte (em pixels)
    """
    try:
        img = PILImage.open(BytesIO(image_bytes))
        
        # Converte para RGB se necessário
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        
        # Converte para array numpy
        img_array = np.array(img)
        
        # Se for RGB, converte para escala de cinza
        if len(img_array.shape) == 3:
            gray = np.mean(img_array, axis=2)
        else:
            gray = img_array
        
        # Define threshold para considerar "branco"
        threshold = 240
        
        # Define a faixa de busca
        start = max(0, target_height - search_range)
        end = min(img.height, target_height + search_range)
        
        # Procura a linha mais "branca" (com menos pixels escuros) na faixa
        best_line = target_height
        min_dark_pixels = float('inf')
        
        for y in range(start, end):
            if y >= img.height:
                break
            # Conta quantos pixels escuros há nesta linha
            dark_pixels = np.sum(gray[y] < threshold)
            
            # Se encontrou uma linha completamente branca, usa ela
            if dark_pixels == 0:
                return y
            
            # Senão, guarda a linha com menos pixels escuros
            if dark_pixels < min_dark_pixels:
                min_dark_pixels = dark_pixels
                best_line = y
        
        return best_line
        
    except Exception as e:
        print(f"Erro ao encontrar ponto de divisão: {e}")
        return target_height


def split_image_at_height(image_bytes: bytes, split_height: int) -> tuple[bytes, bytes]:
    """
    Divide uma imagem em duas partes na altura especificada.
    
    Args:
        image_bytes: Bytes da imagem original
        split_height: Altura onde dividir a imagem
    
    Returns:
        Tupla com (parte_superior_bytes, parte_inferior_bytes)
    """
    try:
        img = PILImage.open(BytesIO(image_bytes))
        
        # Parte superior (do topo até split_height)
        top_img = img.crop((0, 0, img.width, split_height))
        top_output = BytesIO()
        top_img.save(top_output, format='PNG')
        
        # Parte inferior (de split_height até o fim)
        bottom_img = img.crop((0, split_height, img.width, img.height))
        bottom_output = BytesIO()
        bottom_img.save(bottom_output, format='PNG')
        
        return top_output.getvalue(), bottom_output.getvalue()
        
    except Exception as e:
        print(f"Erro ao dividir imagem: {e}")
        # Em caso de erro, retorna a imagem original e uma imagem vazia
        empty_img = PILImage.new('RGB', (1, 1), color='white')
        empty_output = BytesIO()
        empty_img.save(empty_output, format='PNG')
        return image_bytes, empty_output.getvalue()


def crop_whitespace_from_image(image_bytes: bytes, margin: int = 10) -> bytes:
    """
    Remove espaços em branco das bordas de uma imagem, mantendo uma margem mínima.
    
    Args:
        image_bytes: Bytes da imagem original
        margin: Margem mínima a manter em pixels (padrão 10)
    
    Returns:
        Bytes da imagem recortada em formato PNG
    """
    try:
        # Abre a imagem
        img = PILImage.open(BytesIO(image_bytes))
        
        # Converte para RGB se necessário
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        
        # Converte para array numpy para análise
        img_array = np.array(img)
        
        # Se for RGB, converte para escala de cinza para análise
        if len(img_array.shape) == 3:
            gray = np.mean(img_array, axis=2)
        else:
            gray = img_array
        
        # Define threshold para considerar "branco" (255 = branco puro, usamos 240 para tolerância)
        threshold = 240
        
        # Encontra linhas e colunas que não são brancas
        non_white_rows = np.where(np.any(gray < threshold, axis=1))[0]
        non_white_cols = np.where(np.any(gray < threshold, axis=0))[0]
        
        # Se a imagem está completamente branca, retorna uma imagem pequena
        if len(non_white_rows) == 0 or len(non_white_cols) == 0:
            # Cria uma imagem pequena vazia (1x1 px)
            empty_img = PILImage.new('RGB', (1, 1), color='white')
            output = BytesIO()
            empty_img.save(output, format='PNG')
            return output.getvalue()
        
        # Calcula os limites do crop com margem
        top = max(0, non_white_rows[0] - margin)
        bottom = min(img.height, non_white_rows[-1] + margin + 1)
        left = max(0, non_white_cols[0] - margin)
        right = min(img.width, non_white_cols[-1] + margin + 1)
        
        # Recorta a imagem
        cropped_img = img.crop((left, top, right, bottom))
        
        # Converte de volta para bytes
        output = BytesIO()
        cropped_img.save(output, format='PNG')
        
        return output.getvalue()
        
    except Exception as e:
        print(f"Erro ao recortar espaços em branco da imagem: {e}")
        # Em caso de erro, retorna a imagem original
        return image_bytes


def pdf_to_images(pdf_bytes: bytes, dpi: int = 150) -> list[bytes]:
    """
    Converte cada página de um PDF em uma imagem PNG.
    
    Args:
        pdf_bytes: Bytes do arquivo PDF
        dpi: Resolução da imagem (padrão 150 DPI para boa qualidade)
    
    Returns:
        Lista de bytes de imagens PNG, uma para cada página
    """
    images = []
    
    try:
        # Abre o PDF com PyMuPDF
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Converte cada página em imagem
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            
            # Renderiza a página como imagem
            # zoom controla a resolução (2.0 = 144 DPI, ajuste conforme necessário)
            zoom = dpi / 72.0  # 72 é o DPI padrão do PDF
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            # Converte para PNG bytes
            img_bytes = pix.tobytes("png")
            
            # Remove espaços em branco da imagem
            img_bytes_cropped = crop_whitespace_from_image(img_bytes, margin=20)
            images.append(img_bytes_cropped)
            
            # Libera memória
            pix = None
        
        pdf_doc.close()
        
    except Exception as e:
        print(f"Erro ao converter PDF em imagens: {e}")
        raise
    
    return images


def pdf_to_docx_bytes(pdf_bytes: bytes, extract_images: bool = False):
    """
    Converte um PDF em DOCX em memória usando PyMuPDF para texto e pdfplumber para tabelas.
    Mantém melhor a formatação original preservando blocos de texto e detectando tabelas reais.
    Agora também extrai imagens dos PDFs.
    """    
    doc = docx.Document()
    
    # Abre o PDF no PyMuPDF para pegar texto com coordenadas
    try:
        pdf_mupdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        print(f"[DEBUG] pdf_to_docx_bytes: ERRO ao abrir PDF com PyMuPDF: {e}")
        raise
    
    # Lista para armazenar imagens extraídas
    extracted_images = []
    
    # Abre no pdfplumber para pegar tabelas
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf_plumber:   
        for page_num, (page_mupdf, page_plumber) in enumerate(zip(pdf_mupdf, pdf_plumber.pages), start=1):
            
            # 🖼️ Extrai imagens da página usando PyMuPDF
            image_list = page_mupdf.get_images()
            page_images = []
            if extract_images:
                for img_index, img in enumerate(image_list):
                    try:
                        # Obtém dados da imagem
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_mupdf, xref)
                        
                        # Converte para bytes (PNG)
                        if pix.n - pix.alpha < 4:  # Se não é CMYK
                            img_data = pix.tobytes("png")
                            page_images.append((img_data, img_index))
                        else:  # CMYK: converte para RGB primeiro
                            pix_rgb = fitz.Pixmap(fitz.csRGB, pix)
                            img_data = pix_rgb.tobytes("png")
                            page_images.append((img_data, img_index))
                            pix_rgb = None
                        
                        pix = None
                    except Exception as e:
                        print(f"Erro ao extrair imagem {img_index} da página {page_num}: {e}")
            
            # 1️⃣ Extrai blocos de texto (mantém ordem visual)
            blocks = page_mupdf.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)
            blocks_sorted = sorted(blocks, key=lambda b: (b[1], b[0]))  # ordena por posição (y, x)
            
            # 2️⃣ Extrai tabelas com pdfplumber
            tables = page_plumber.extract_tables()
            table_bboxes = []
            
            # Pega as coordenadas das tabelas para evitar duplicação de conteúdo
            for table_obj in page_plumber.find_tables():
                if table_obj.bbox:
                    table_bboxes.append(table_obj.bbox)
            
            # 3️⃣ Processa blocos de texto, evitando áreas de tabelas
            # Primeiro, filtra blocos fora de tabelas e consolida parágrafos baseado na posição Y
            text_blocks = []
            for block in blocks_sorted:
                x0, y0, x1, y1, text, block_no, block_type = block
                
                # Verifica se o bloco está dentro de uma área de tabela
                is_in_table = False
                for table_bbox in table_bboxes:
                    tx0, ty0, tx1, ty1 = table_bbox
                    # Verifica sobreposição significativa
                    if not (x1 < tx0 or x0 > tx1 or y1 < ty0 or y0 > ty1):
                        overlap_area = (min(x1, tx1) - max(x0, tx0)) * (min(y1, ty1) - max(y0, ty0))
                        block_area = (x1 - x0) * (y1 - y0)
                        if block_area > 0 and overlap_area / block_area > 0.3:  # 30% de sobreposição
                            is_in_table = True
                            break
                
                if not is_in_table and text.strip():
                    # Limpa o texto removendo informações de autenticação
                    clean_text = re.sub(
                        r"A autenticidade deste documento pode ser conferida.*?(https?://\S+)",
                        "", text, flags=re.DOTALL
                    )
                    clean_text = re.sub(
                        r"^\s*\d{6,}-[A-Za-z0-9]{4,}\s*$",
                        "", clean_text, flags=re.MULTILINE
                    )
                    
                    if clean_text.strip():
                        text_blocks.append((y0, y1, clean_text.strip()))
            
            # Consolida blocos em parágrafos baseado na posição Y
            consolidated_paragraphs = []
            current_paragraph = ""
            last_y_end = None
            
            for y0, y1, text in text_blocks:
                # Se é o primeiro bloco ou há uma grande distância vertical (nova seção)
                if last_y_end is None or (y0 - last_y_end) > 15:  # 15 pontos de distância
                    # Finaliza parágrafo anterior se existir
                    if current_paragraph.strip():
                        consolidated_paragraphs.append(current_paragraph.strip())
                    current_paragraph = text
                else:
                    # Consolida com parágrafo atual
                    if current_paragraph and not current_paragraph.endswith(' '):
                        current_paragraph += " "
                    current_paragraph += text
                
                last_y_end = y1
            
            # Adiciona último parágrafo
            if current_paragraph.strip():
                consolidated_paragraphs.append(current_paragraph.strip())
            
            # Processa parágrafos consolidados
            for para in consolidated_paragraphs:
                if para.strip():
                    # Detecta padrão nome + cargo na mesma linha
                    match = re.match(r'^([A-ZÁÉÍÓÚÂÊÔÃÕÇ ]{5,})\s+([A-Z][a-z].+)$', para.strip())
                    if match:
                        nome = match.group(1).strip()
                        cargo = match.group(2).strip()
                        doc.add_paragraph(nome)
                        doc.add_paragraph(cargo)
                    else:
                        doc.add_paragraph(para)
            
            # 🖼️ Adiciona imagens extraídas da página (caso não foram incluídas no texto)
            for img_data, img_index in page_images:
                # Adiciona um parágrafo vazio antes da imagem para separação
                doc.add_paragraph()
                # Salva referência para processamento posterior
                extracted_images.append(img_data)
            
            # 4️⃣ Processa tabelas reais
            for table in tables:
                if table and len(table) > 1 and len(table[0]) > 1:
                    # Verifica se a tabela tem conteúdo válido
                    has_content = any(
                        any(cell and cell.strip() for cell in row) 
                        for row in table
                    )
                    
                    if has_content:
                        # Aplica consolidação de linhas fragmentadas
                        consolidated_table = _consolidate_table_rows(table)
                        
                        # Consolida cabeçalhos fragmentados (corrige vazamento de conteúdo)
                        headers_consolidated_table = _consolidate_fragmented_headers(consolidated_table)
                        
                        # Remove linhas de título mesclado (primeira célula com conteúdo longo, demais vazias)
                        no_merged_headers_table = _remove_merged_header_rows(headers_consolidated_table)
                        
                        # Remove colunas duplicadas e vazias, preservando conteúdo
                        cleaned_table = _remove_duplicate_and_empty_columns(no_merged_headers_table)
                        
                        rows = len(cleaned_table)
                        cols = len(cleaned_table[0]) if rows > 0 else 0
                        
                        # Cria tabela no DOCX
                        table_docx = doc.add_table(rows=rows, cols=cols)
                        
                        for i, row in enumerate(cleaned_table):
                            for j, cell in enumerate(row):
                                if j < len(table_docx.rows[i].cells):
                                    table_docx.rows[i].cells[j].text = cell if cell else ""
                        
                        # Adiciona espaço após a tabela
                        doc.add_paragraph()
            
            # Adiciona quebra de página se não for a última página
            if page_num < len(pdf_mupdf):
                doc.add_page_break()
    
    # Fecha o PDF do PyMuPDF
    pdf_mupdf.close()
    
    # Salva o documento em BytesIO
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    
    # Retorna tupla com (docx_bytes, imagens_extraídas)
    # Se não extraiu imagens, retorna apenas o BytesIO para compatibilidade
    if extract_images:
        return (out, extracted_images)
    else:
        return out


def consolidate_cross_page_tables(elements):
    """
    Consolida tabelas que foram divididas por mudança de página.
    Detecta tabelas com cabeçalhos idênticos e as mescla.
    """
    if len(elements) < 2:
        return elements
    
    # Encontra grupos de tabelas com cabeçalhos idênticos
    table_groups = []
    current_group = []
    
    for i, (element_type, element) in enumerate(elements):
        if element_type == 'table':
            rows = len(element.rows) if hasattr(element, 'rows') else 0
            
            if rows > 0:
                # Extrai cabeçalho normalizado
                current_header = [cell.text.strip().lower() for cell in element.rows[0].cells]
                
                # Se é a primeira tabela ou tem cabeçalho diferente
                if not current_group:
                    current_group = [(i, element, current_header)]
                else:
                    # Compara com cabeçalho do último elemento do grupo atual
                    last_header = current_group[-1][2]
                    
                    if current_header == last_header:
                        current_group.append((i, element, current_header))
                    else:
                        # Cabeçalho diferente - finaliza grupo anterior e inicia novo
                        if len(current_group) > 1:
                            table_groups.append(current_group)
                        current_group = [(i, element, current_header)]
    
    # Adiciona último grupo se tiver mais de uma tabela
    if len(current_group) > 1:
        table_groups.append(current_group)
    
    # Se não há grupos para mesclar, retorna elementos originais
    if not table_groups:
        return elements
    
    # Constrói novo array de elementos com tabelas mescladas
    new_elements = []
    processed_positions = set()
    
    # Para cada grupo, substitui as tabelas por uma única tabela mesclada
    for group in table_groups:
        merged_table = merge_table_group(group)
        
        if merged_table:
            # Marca todas as posições do grupo como processadas
            first_pos = group[0][0]
            for pos, _, _ in group:
                processed_positions.add(pos)
    
    # Reconstrói lista de elementos
    for i, element in enumerate(elements):
        if i not in processed_positions:
            new_elements.append(element)
        elif any(group[0][0] == i for group in table_groups):
            # É a primeira posição de um grupo - adiciona tabela mesclada
            for group in table_groups:
                if group[0][0] == i:
                    merged_table = merge_table_group(group)
                    if merged_table:
                        new_elements.append(('table', merged_table))
                    break
    
    return new_elements

def merge_table_group(table_group):
    """Mescla um grupo de tabelas com cabeçalhos idênticos"""
    if len(table_group) < 2:
        return None
    
    # Calcula dimensões da nova tabela
    total_data_rows = 0
    cols = 0
    
    for pos, table, header in table_group:
        rows = len(table.rows) if hasattr(table, 'rows') else 0
        table_cols = len(table.rows[0].cells) if rows > 0 else 0
        
        if rows > 1:  # Ignora apenas cabeçalho
            total_data_rows += rows - 1  # -1 para não contar cabeçalho
        
        cols = max(cols, table_cols)
    
    # Total: cabeçalho + todas as linhas de dados
    total_rows = 1 + total_data_rows
    
    # Cria documento temporário para nova tabela
    temp_doc = docx.Document()
    merged_table = temp_doc.add_table(rows=total_rows, cols=cols)
    
    # Adiciona cabeçalho (pega da primeira tabela)
    first_table = table_group[0][1]
    if hasattr(first_table, 'rows') and len(first_table.rows) > 0:
        for j, cell in enumerate(first_table.rows[0].cells):
            if j < len(merged_table.rows[0].cells):
                merged_table.rows[0].cells[j].text = cell.text.strip()
    
    # Adiciona dados de todas as tabelas
    current_row = 1
    for pos, table, header in table_group:
        if hasattr(table, 'rows'):
            # Copia todas as linhas exceto cabeçalho
            for i in range(1, len(table.rows)):
                if current_row < len(merged_table.rows):
                    for j, cell in enumerate(table.rows[i].cells):
                        if j < len(merged_table.rows[current_row].cells):
                            merged_table.rows[current_row].cells[j].text = cell.text.strip()
                    current_row += 1
    
    return merged_table

def extract_elements_from_docx(docx_data,image:bool=False):
    docx_file = BytesIO(docx_data)
    #Passa os BytesIO pra função que retorna um Document 
    try:
        doc = docx.Document(docx_file)
    except Exception as e:
        print(f"[DEBUG] ERRO ao criar documento DOCX: {e}")
        return []
    
    elements = []
    
    #Passa o objeto Document e pra yielda um objeto tipo Table ou Paragraph
    block_count = 0
    for block in iter_block_items(doc):
        block_count += 1
        #Verifica se qual o tipo do block e adiciona uma tupla (tipo:str,block:(Table ou Paragraph))
        if isinstance(block, docx.table.Table):
            elements.append(('table', block))
        elif isinstance(block, docx.text.paragraph.Paragraph):
            text_preview = block.text[:50] if block.text else "(vazio)"
            elements.append(('paragraph', block))
        else:
            print(f"[DEBUG] Block {block_count}: Tipo desconhecido - {type(block)}")
    
    
    # Se devemos extrair imagens, adiciona TODAS as imagens do documento usando método mais robusto
    if image:
        all_docx_images = extract_images_from_docx(docx_data)
        for i, image_data in enumerate(all_docx_images):
            elements.append(('image', image_data))
    
    # Aplica consolidação de tabelas divididas por mudança de página
    elements = consolidate_cross_page_tables(elements)    
    return elements
        

def iter_block_items(parent):
    from docx.oxml import CT_P, CT_Tbl
    from docx.oxml.ns import qn
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

def extract_images_from_docx(docx_data):
    """
    Extrai imagens de um documento DOCX.
    Retorna uma lista de dados binários das imagens encontradas.
    """
    docx_file = BytesIO(docx_data)
    doc = docx.Document(docx_file)
    images = []
    
    try:
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                image_part = rel.target_part
                image_data = image_part.blob
                images.append(image_data)
    except Exception as e:
        print(f"Erro ao extrair imagens do DOCX: {e}")
    
    return images

def extract_table_data(table, cell_style):
    long_cell_style = ParagraphStyle(
        'CellStyle',
        fontName="Calibri",
        fontSize=5,
        leading=10,
        alignment=TA_JUSTIFY,
        wordWrap='CJK'  
    )
    
    # Extrai dados brutos da tabela DOCX
    raw_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = '\n'.join([para.text for para in cell.paragraphs])
            row_data.append(cell_text)
        raw_data.append(row_data)
    
    # Aplica correções: primeiro consolida linhas fragmentadas, depois cabeçalhos fragmentados, depois remove títulos mesclados, depois remove colunas duplicadas
    if raw_data:
        consolidated_data = _consolidate_table_rows(raw_data)
        headers_consolidated_data = _consolidate_fragmented_headers(consolidated_data)
        no_merged_headers_data = _remove_merged_header_rows(headers_consolidated_data)
        cleaned_data = _remove_duplicate_and_empty_columns(no_merged_headers_data)
    else:
        cleaned_data = raw_data
    
    # Converte para objetos Paragraph do ReportLab
    data = []
    for row in cleaned_data:
        row_data = []
        for cell_text in row:
            if len(cell_text) > 50:
                long_cell = Paragraph(cell_text, long_cell_style)
                row_data.append(long_cell)
            else:    
                cell_para = Paragraph(cell_text, cell_style)
                row_data.append(cell_para)
        data.append(row_data)
    
    return data

def is_all_upper(s):
    # Normaliza a string para garantir consistência nos caracteres
    s_normalized = unicodedata.normalize('NFC', s)
    # Itera sobre os caracteres da string
    for c in s_normalized:
        # Verifica se é uma letra que pode ter distinção entre maiúsculo e minúsculo
        if c.isalpha() and c.lower() != c.upper():
            # Verifica se a letra não é maiúscula
            if not c.isupper():
                return False
    # Retorna True se todas as letras relevantes forem maiúsculas
    return True

pdfmetrics.registerFont(TTFont('Calibri',os.path.join(FONTS_DIR,'Calibri.ttf')))
pdfmetrics.registerFont(TTFont('Calibri-Bold', os.path.join(FONTS_DIR,'Calibri-Bold.ttf')))
pdfmetrics.registerFont(TTFont('Impact', os.path.join(FONTS_DIR,'Impact.ttf')))

def create_pdf(documents, context):
    """
    Cria PDF do diário oficial
    
    Args:
        documents: Lista de objetos DocumentProcessed
        context: Contexto com dados do estabelecimento e configurações
    """
    documentos_por_orgao = defaultdict(list)
    documentos_imagem = []  # Lista separada para documentos tipo='Imagem' (extração de imagens do PDF)
    for doc in documents:
        orgao = doc.orgao
        if orgao.is_estabelecimento:
            documentos_por_orgao[orgao.nome].append(doc)

    for document in documents:
        # Apenas documentos tipo='Imagem' vão para seção separada
        # Documentos com force_scan mantêm sua ordem normal por órgão
        if document.tipo == 'Imagem':
            documentos_imagem.append(document)
        else:
            if document.orgao.is_estabelecimento:
                continue  # Já processado na seção de estabelecimento
            orgao_nome = document.orgao.nome
            documentos_por_orgao[orgao_nome].append(document)

    buffer = BytesIO()

    doc = BaseDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()

    width, height = A4
    margin = 1 * cm
    gutter = 0.8 * cm
    column_width = (width - 2 * margin - gutter) / 2
    usable_height = height - 2.2 * margin
    header_height = 5 * cm
    
    # Espaço reservado para container de assinatura (será adicionado depois)
    assinatura_space = 4 * cm  # 3.5cm do container + 0.5cm de espaço

    frame1_first = Frame(
        margin,
        margin + 6.5 * cm,
        column_width,
        usable_height - header_height - 5 * cm - assinatura_space,  # Reduz altura para reservar espaço
        id='col1_first',
    )
    frame2_first = Frame(
        margin + column_width + gutter,
        margin + 1.5*cm,
        column_width,
        usable_height - header_height,
        id='col2_first'
    )

    frame1_later = Frame(
        margin,
        margin - 0.3 * cm,
        column_width,
        usable_height,
        id='col1_later'
    )
    frame2_later = Frame(
        margin + column_width + gutter,
        margin - 0.3 * cm,
        column_width,
        usable_height,
        id='col2_later',
    )

    first_page_template = PageTemplate(
        id='FirstPage',
        frames=[frame1_first, frame2_first],
        onPage=create_header_template(context),
        onPageEnd=create_container_responsavel(context),
        autoNextPageTemplate='LaterPages'
    )

    later_page_template = PageTemplate(
        id='LaterPages',
        frames=[frame1_later, frame2_later],
        onPage=create_header_template_later(context),
        onPageEnd=footer_template,
        autoNextPageTemplate='LaterPages'
    )

    doc.addPageTemplates([first_page_template, later_page_template])

    title_style = ParagraphStyle(
        'TitleStyle',
        fontName="Calibri-Bold",
        fontSize=9,
        leading=10,
        spaceAfter=0,
        alignment=TA_CENTER
    )

    left_style = ParagraphStyle(
        'LeftStyle',
        fontName="Calibri",
        fontSize=9,
        leading=10,
        spaceAfter=0,
        alignment=TA_LEFT
    )

    section_title_style = ParagraphStyle(
        'SectionTitleStyle',
        fontName="Calibri-Bold",
        fontSize=10,
        leading=12,
        spaceBefore=12,
        spaceAfter=6,
        alignment=TA_LEFT
    )

    centered_style = ParagraphStyle(
        'CenteredStyle',
        fontName="Calibri",
        fontSize=8,
        leading=10,
        alignment=TA_CENTER
    )

    normal_style = ParagraphStyle(
        'NormalStyle',
        fontName="Calibri",
        fontSize=8,
        leading=10,
        firstLineIndent=20,
        alignment=TA_JUSTIFY
    )

    name_style = ParagraphStyle(
        'NameStyle',
        fontName="Calibri",
        fontSize=9,
        leading=5,
        spaceAfter=0,
        alignment=TA_CENTER
    )

    cell_style = ParagraphStyle(
        'CellStyle',
        fontName="Calibri",
        fontSize=5,
        leading=10,
        alignment=TA_CENTER,
        wordWrap='CJK'  
    )

    table_style = TableStyle([
                                ('FONTNAME', (0, 0), (-1, -1), 'Calibri'),
                                ('FONTSIZE', (0, 0), (-1, -1), 8),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('LEFTPADDING', (0, 0), (-1, -1), 2),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                                ('TOPPADDING', (0, 0), (-1, -1), 2),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                                ('WORDWRAP', (0, 0), (-1, -1), 'CJK'), 
                                ('BOX', (0, 0), (-1, -1), 1, 'black'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), 
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
                        ])


    story = []


    # Primeiro processa documentos normais por poder
    for orgao in documentos_por_orgao:
        section_title_str = str(orgao).upper()
        section_title = SectionTitle(section_title_str, column_width - gutter, 1.0 * cm, offset_x=-0.3 * cm)
        story.append(section_title)
        story.append(Spacer(1, 12))
        for doc_item in sorted(documentos_por_orgao[orgao]):
            try:
                first_iteration = True
                elements = None
                
                # Verifica se deve forçar scan do PDF para imagem
                force_scan = getattr(doc_item, 'force_scan', False)
                
                # Verifica se é PDF pela extensão do filename OU pelos magic bytes do arquivo
                is_pdf = doc_item.filename.lower().endswith('.pdf')
                if not is_pdf and len(doc_item.data) >= 4:
                    # Verifica magic bytes do PDF (%PDF)
                    is_pdf = doc_item.data[:4] == b'%PDF'
                
                # Verifica se é DOCX pela extensão OU pelos magic bytes
                is_docx = doc_item.filename.lower().endswith('.docx') or doc_item.filename.lower().endswith('.doc')
                if not is_docx and len(doc_item.data) >= 4:
                    # DOCX são arquivos ZIP, então começam com 'PK' (magic bytes do ZIP)
                    # Verificação mais robusta: PK\x03\x04 para DOCX/ZIP
                    is_docx = doc_item.data[:2] == b'PK'

                if force_scan and is_pdf:
                    # Converte PDF para imagens e adiciona como elementos de imagem
                    pdf_images = pdf_to_images(doc_item.data)
                    elements = []
                    # Adiciona cada página como imagem
                    for idx, img_bytes in enumerate(pdf_images):
                        elements.append(('image', img_bytes))
                elif is_docx:
                    elements = extract_elements_from_docx(doc_item.data)
                elif is_pdf:
                    fake_docx_bio = pdf_to_docx_bytes(doc_item.data)
                    fake_docx_bytes = fake_docx_bio.getvalue()
                    elements = extract_elements_from_docx(fake_docx_bytes)
                else:
                    elements = None
                
                if elements is None:
                    continue  # Pula este documento e continua com o próximo
                                                        
                element_count = 0
                for element_type, element in elements:
                    element_count += 1
                    
                    if element_type == 'paragraph':
                        paragraph = element.text
                        stripped_paragraph = ''.join(paragraph.split(' '))
                        if paragraph and first_iteration:
                            story.append(Paragraph(paragraph, title_style))
                            first_iteration = False
                            continue
                        if len(stripped_paragraph) <= 50 and is_name(paragraph):
                            story.append(Paragraph(paragraph.upper(), name_style))
                        elif len(stripped_paragraph) <= 70 and is_all_upper(stripped_paragraph):
                            story.append(Paragraph(paragraph, title_style))
                        elif (stripped_paragraph.startswith("-") or stripped_paragraph.endswith(":") and len(paragraph) <= 50):
                            story.append(Paragraph(paragraph, left_style))
                        elif len(stripped_paragraph) <= 50:
                            story.append(Paragraph(paragraph, centered_style))
                        else:
                            story.append(Paragraph(paragraph, normal_style))
                        story.append(Spacer(1, 4))

                    elif element_type == 'image':
                        # Processa elementos de imagem (incluindo PDFs escaneados)
                        # Agora com divisão automática se não couber na coluna!
                        try:
                            document_image = create_splittable_image(
                                element,
                                max_width=column_width - 1.15*cm,
                                max_height=usable_height - 2*cm,  # Altura máxima considerando margens
                                alignment='center'
                            )
                            story.append(document_image)
                            story.append(Spacer(1, 8))
                        except Exception as e:
                            print(f"Erro ao processar imagem no documento {doc_item.filename}: {e}")
                            story.append(Spacer(1, 4))

                    elif element_type == 'table':
                        table_data = extract_table_data(element, cell_style)
                        num_cols = len(table_data[0]) if table_data else 0
                        if num_cols > 0:
                            total_table_width = column_width
                            min_width = 1 * cm
                            max_width = 4 * cm
                            from reportlab.pdfbase.pdfmetrics import stringWidth
                            column_widths_raw = []
                            for col in zip(*table_data):
                                max_column_width = max(stringWidth(str(cell.text), 'Calibri', 5) for cell in col)
                                column_widths_raw.append(max(min(max_column_width, max_width), min_width))
                            total_raw_width = sum(column_widths_raw)
                            col_widths = [
                                total_table_width * (raw_width / total_raw_width) for raw_width in column_widths_raw
                            ]
                            reportlab_table = Table(table_data, colWidths=col_widths, splitInRow=1)
                            table_width, table_height = reportlab_table.wrap(total_table_width, usable_height)
                            from reportlab.platypus import FrameBreak
                            if table_height > usable_height:
                                story.append(FrameBreak())
                            reportlab_table.setStyle(table_style)
                            story.append(reportlab_table)
                            story.append(Spacer(1, 10))

                story.append(Spacer(1, 15))
            except Exception as e:
                raise DocumentReadError(doc_item.filename, e)

    # Depois processa documentos de imagem (sempre no final do diário)
    if documentos_imagem:
        # Adiciona seção para imagens se houver documentos deste tipo
        section_title = SectionTitle("IMAGENS", column_width - gutter, 1.0 * cm, offset_x=-0.3 * cm)
        story.append(section_title)
        story.append(Spacer(1, 12))
        
        for doc_item in sorted(documentos_imagem):
            try:
                first_iteration = True
                elements = None
                extracted_pdf_images = []  # Para armazenar imagens de PDFs
                
                # Verifica se deve forçar scan do PDF para imagem
                force_scan = getattr(doc_item, 'force_scan', False)
                
                # Verifica tipo de arquivo
                is_pdf = doc_item.filename.lower().endswith('.pdf')
                if not is_pdf and len(doc_item.data) >= 4:
                    is_pdf = doc_item.data[:4] == b'%PDF'
                
                # Verifica se é DOCX pela extensão OU pelos magic bytes
                is_docx = doc_item.filename.lower().endswith('.docx') or doc_item.filename.lower().endswith('.doc')
                if not is_docx and len(doc_item.data) >= 4:
                    # DOCX são arquivos ZIP, então começam com 'PK'
                    is_docx = doc_item.data[:2] == b'PK'
                
                if force_scan and is_pdf:
                    # Converte PDF para imagens e adiciona como elementos de imagem
                    pdf_images = pdf_to_images(doc_item.data)
                    elements = []
                    # Adiciona cada página como imagem
                    for img_bytes in pdf_images:
                        elements.append(('image', img_bytes))
                elif is_docx:
                    elements = extract_elements_from_docx(doc_item.data, image=True)
                elif is_pdf:
                    # Converte PDF para DOCX com extração de imagens
                    result = pdf_to_docx_bytes(doc_item.data, extract_images=True)
                    
                    # Verifica se retornou tupla (docx, imagens) ou apenas docx
                    if isinstance(result, tuple):
                        fake_docx_bio, extracted_pdf_images = result
                        fake_docx_bytes = fake_docx_bio.getvalue()
                    else:
                        fake_docx_bytes = result.getvalue()
                        extracted_pdf_images = []
                    
                    elements = extract_elements_from_docx(fake_docx_bytes, image=True)
                    
                    # Adiciona imagens extraídas do PDF aos elementos
                    for img_data in extracted_pdf_images:
                        elements.append(('image', img_data))
                else:
                    print(f"[DEBUG] AVISO: Tipo de arquivo não reconhecido para documento IMAGEM!")
                    elements = []
                
                if elements is None:
                    print(f"[DEBUG] ERRO: elements é None, pulando documento")
                    continue
                
                if not elements:
                    print(f"[DEBUG] AVISO: Lista de elements está vazia!")
                                
                for element_type, element in elements:
                    if element_type == 'paragraph':
                        paragraph = element.text
                        stripped_paragraph = ''.join(paragraph.split(' '))
                        if paragraph and first_iteration:
                            story.append(Paragraph(paragraph, title_style))
                            first_iteration = False
                            continue
                        if len(stripped_paragraph) <= 50 and is_name(paragraph):
                            story.append(Paragraph(paragraph.upper(), name_style))
                        elif len(stripped_paragraph) <= 70 and is_all_upper(stripped_paragraph):
                            story.append(Paragraph(paragraph, title_style))
                        elif (stripped_paragraph.startswith("-") or stripped_paragraph.endswith(":") and len(paragraph) <= 50):
                            story.append(Paragraph(paragraph, left_style))
                        elif len(stripped_paragraph) <= 50:
                            story.append(Paragraph(paragraph, centered_style))
                        else:
                            story.append(Paragraph(paragraph, normal_style))
                        story.append(Spacer(1, 4))

                    elif element_type == 'image':
                        # Processa elementos de imagem
                        # Agora com divisão automática se não couber na coluna!
                        try:
                            # Cria o elemento de imagem com largura máxima da coluna
                            document_image = create_splittable_image(
                                element, 
                                max_width=column_width - 3*cm,  # Deixa uma margem
                                max_height=8*cm,  # Altura máxima razoável
                                alignment='center'
                            )
                            story.append(document_image)
                            story.append(Spacer(1, 8))  # Espaço após a imagem
                        except Exception as e:
                            print(f"Erro ao processar imagem no documento {doc_item.filename}: {e}")
                            # Adiciona espaço em caso de erro
                            story.append(Spacer(1, 4))

                    elif element_type == 'table':
                        table_data = extract_table_data(element, cell_style)
                        num_cols = len(table_data[0]) if table_data else 0
                        if num_cols > 0:
                            total_table_width = column_width
                            min_width = 1 * cm
                            max_width = 4 * cm
                            from reportlab.pdfbase.pdfmetrics import stringWidth
                            column_widths_raw = []
                            for col in zip(*table_data):
                                max_column_width = max(stringWidth(str(cell.text), 'Calibri', 5) for cell in col)
                                column_widths_raw.append(max(min(max_column_width, max_width), min_width))
                            total_raw_width = sum(column_widths_raw)
                            col_widths = [
                                total_table_width * (raw_width / total_raw_width) for raw_width in column_widths_raw
                            ]
                            reportlab_table = Table(table_data, colWidths=col_widths, splitInRow=1)
                            table_width, table_height = reportlab_table.wrap(total_table_width, usable_height)
                            from reportlab.platypus import FrameBreak
                            if table_height > usable_height:
                                story.append(FrameBreak())
                            reportlab_table.setStyle(table_style)
                            story.append(reportlab_table)
                            story.append(Spacer(1, 10))

                story.append(Spacer(1, 15))
            except Exception as e:
                raise DocumentReadError(doc_item.filename, e)

    doc.build(story)
    pdf_data = buffer.getvalue()
    buffer.seek(0)
    buffer.close()
    return pdf_data


from PyPDF2 import PdfReader, PdfWriter
import qrcode

SIGNATURE_PRIMARY = colors.HexColor("#03A64A")
SIGNATURE_SECONDARY = colors.HexColor("#0B5E8E")
SIGNATURE_TEXT = colors.HexColor("#27332E")
SIGNATURE_MUTED = colors.HexColor("#5F6F66")
SIGNATURE_BORDER = colors.HexColor("#CDD9D2")
SIGNATURE_BACKGROUND = colors.HexColor("#FBFDFC")
SIGNATURE_SOFT_BACKGROUND = colors.HexColor("#F1F7F3")
SIGNATURE_SHADOW = colors.HexColor("#DDE7E1")


def _safe_signature_text(value):
    if value is None:
        return ""
    return escape(str(value).strip())


def _truncate_middle(value, max_length=46):
    text = str(value or "").strip()
    if len(text) <= max_length:
        return text
    if max_length <= 3:
        return text[:max_length]

    visible_length = max_length - 3
    left_length = visible_length // 2
    right_length = visible_length - left_length
    return f"{text[:left_length]}...{text[-right_length:]}"


def _format_signature_datetime(raw_value):
    if not raw_value:
        return ""

    original_value = str(raw_value).strip()
    normalized_value = original_value[2:] if original_value.startswith("D:") else original_value
    normalized_value = normalized_value.replace("'", "")

    if len(normalized_value) < 14 or not normalized_value[:14].isdigit():
        return original_value

    try:
        parsed_datetime = datetime.strptime(normalized_value[:14], "%Y%m%d%H%M%S")
    except ValueError:
        return original_value

    utc_suffix = ""
    timezone_offset = normalized_value[14:]
    if timezone_offset.startswith("Z"):
        utc_suffix = " UTC"
    elif (
        len(timezone_offset) >= 5
        and timezone_offset[0] in "+-"
        and timezone_offset[1:5].isdigit()
    ):
        utc_suffix = f" UTC{timezone_offset[:3]}:{timezone_offset[3:5]}"

    return f"{parsed_datetime.strftime('%d/%m/%Y %H:%M')}{utc_suffix}"


def _format_validation_display(url_qrcode):
    if not url_qrcode:
        return ""

    parsed_url = urlparse(str(url_qrcode).strip())
    if parsed_url.scheme and parsed_url.netloc:
        last_segment = parsed_url.path.rstrip("/").split("/")[-1]
        if last_segment:
            compact_url = f"{parsed_url.netloc}/.../{last_segment}"
        else:
            compact_url = parsed_url.netloc
        return _truncate_middle(compact_url, 44)

    return _truncate_middle(url_qrcode, 44)


def _add_flowables_in_frame(canv, x, y, width, height, flowables):
    frame = Frame(
        x,
        y,
        width,
        height,
        showBoundary=0,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    frame.addFromList(list(flowables), canv)


def _merge_overlay_on_first_page(pdf_bytes, overlay_buffer):
    # PyPDF2's merge_page does not properly transfer Image XObjects (only vector
    # graphics survive), so the QR code image ends up blank. Use PyMuPDF instead,
    # which correctly embeds all resources including images.
    overlay_buffer.seek(0)
    overlay_data = overlay_buffer.read()

    main_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    overlay_doc = fitz.open(stream=overlay_data, filetype="pdf")

    first_page = main_doc[0]
    first_page.show_pdf_page(first_page.rect, overlay_doc, 0, overlay=True)

    result = main_doc.tobytes()

    overlay_doc.close()
    main_doc.close()

    return result


def adicionar_dados_assinatura_footer(diario_bytes, emissor,dados_assinatura):
    assinatura = '''
            DOCUMENTO ASSINADO DIGITALMENTE CONFORME MP N° 2.200-2 DE 24/08/2001, QUE INSTITUI A INFRAESTRUTURA DE CHAVES PÚBLICAS BRASILEIRA - ICP-BRASIL.
        '''

    dados_assinatura = dados_assinatura or {}
    emissor = emissor or dados_assinatura.get('proprietario') or ''

    buffer_texto = BytesIO()
    canv = canvas.Canvas(buffer_texto, pagesize=A4)

    ass_style = ParagraphStyle(
        'LeftStyle',
        fontName="Calibri-Bold",
        fontSize=6,
        leading=6,
        spaceAfter=0,
        alignment=TA_LEFT
    )

    emissor_style = ParagraphStyle(
        'LeftStyle',
        fontName="Calibri",
        fontSize=7,
        leading=6,
        spaceAfter=0,
        alignment=TA_LEFT
    )

    data_cert_style = ParagraphStyle(
        'LeftStyle',
        fontName="Calibri",
        fontSize=5,
        leading=5,
        spaceAfter=0,
        alignment=TA_LEFT
    )

    assinatura_paragraph = Paragraph(assinatura, ass_style)
    emissor_paragraph = Paragraph(emissor, emissor_style)
    dados_certificado = Paragraph(
        f'Assinado Digitalmente por {emissor} Dados: {dados_assinatura.get("signingdate")}',
        data_cert_style,
    )

    x, y = 0.8 * cm, 0 * cm
    largura = 10 * cm
    altura = 1 * cm
    frame = Frame(x, y, largura, altura)
    frame.addFromList([assinatura_paragraph], canv)

    x2 = x + largura
    frame2 = Frame(x2, y, 3.5 * cm, altura)
    frame2.addFromList([emissor_paragraph], canv)

    frame3 = Frame(x2 + 3.5 * cm, y, 3.5 * cm, altura)
    frame3.addFromList([dados_certificado], canv)

    canv.save()
    return _merge_overlay_on_first_page(diario_bytes, buffer_texto)


def adicionar_qrcode_assinatura_final(diario_bytes, url_qrcode, emissor=None, dados_assinatura=None):
    """Adiciona container com QR code na primeira página, coluna esquerda (espaço reservado)."""
    width, height = A4
    margin = 1 * cm
    gutter = 0.8 * cm
    column_width = (width - 2 * margin - gutter) / 2

    # Gerar QR code
    qr_img = qrcode.make(url_qrcode)
    qr_buffer = BytesIO()
    qr_img.save(qr_buffer, 'PNG')
    qr_buffer.seek(0)

    dados_assinatura = dados_assinatura or {}
    emissor_display = _truncate_middle(
        emissor or dados_assinatura.get('proprietario') or '',
        48,
    )
    data_assinatura = _format_signature_datetime(dados_assinatura.get('signingdate'))
    codigo_referencia = _truncate_middle(dados_assinatura.get('cod_ref', ''), 28)
    validacao_display = _format_validation_display(url_qrcode)

    buffer_overlay = BytesIO()
    canv = canvas.Canvas(buffer_overlay, pagesize=A4)

    container_height = 3.55 * cm
    container_width = column_width
    container_x = (margin/2) + (column_width - container_width) / 2
    container_y = height - 4.75 * cm - container_height

    canv.setFillColor(SIGNATURE_BACKGROUND)
    canv.setStrokeColor(SIGNATURE_BORDER)
    canv.setLineWidth(0.8)
    canv.rect(container_x, container_y, container_width, container_height, stroke=1, fill=1)

    qr_box_x = container_x + 0.52 * cm
    qr_box_y = container_y + 0.55 * cm
    qr_box_size = 2.28 * cm
    canv.setFillColor(colors.white)
    canv.rect(qr_box_x, qr_box_y, qr_box_size, qr_box_size, stroke=0, fill=1)
    canv.setStrokeColor(SIGNATURE_BORDER)
    canv.setLineWidth(0.6)
    canv.rect(qr_box_x, qr_box_y, qr_box_size, qr_box_size, stroke=1, fill=0)

    image = ImageReader(qr_buffer)
    qr_padding = 0.12 * cm
    canv.drawImage(
        image,
        qr_box_x + qr_padding,
        qr_box_y + qr_padding,
        width=qr_box_size - (2 * qr_padding),
        height=qr_box_size - (2 * qr_padding),
        preserveAspectRatio=True,
        mask='auto',
    )

    divider_x = qr_box_x + qr_box_size + 0.38 * cm
    canv.setStrokeColor(SIGNATURE_BORDER)
    canv.setLineWidth(0.5)
    canv.line(divider_x, container_y + 0.42 * cm, divider_x, container_y + container_height - 0.42 * cm)

    body_style = ParagraphStyle(
        'SignatureCardBody',
        fontName="Calibri",
        fontSize=5.5,
        leading=6.5,
        textColor=SIGNATURE_TEXT,
        alignment=TA_JUSTIFY,
    )
    code_style = ParagraphStyle(
        'SignatureCode',
        fontName="Calibri-Bold",
        fontSize=7,
        leading=8.5,
        textColor=SIGNATURE_TEXT,
        alignment=TA_LEFT,
    )

    flowables = [
        Paragraph(
            "Documento assinado digitalmente conforme MP N° 2.200-2 de 24/08/2001, da Infraestrutura de "
            "Chaves Públicas Brasileira - ICP-Brasil.",
            body_style,
        ),
        Spacer(1, 0.1 * cm),
        Paragraph(
            "A autenticidade deste documento pode ser conferida por meio do QRCode.",
            body_style,
        ),
        Spacer(1, 0.12 * cm),
    ]

    if codigo_referencia:
        flowables.append(
            Paragraph(
                f"<b>Código de Validação:</b> {_safe_signature_text(codigo_referencia)}",
                code_style,
            )
        )

    text_x = divider_x + 0.28 * cm
    text_y = container_y + 0.45 * cm
    text_width = (container_x + container_width) - text_x - 0.42 * cm
    text_height = container_height - 1 * cm
    _add_flowables_in_frame(canv, text_x, text_y, text_width, text_height, flowables)

    canv.save()
    return _merge_overlay_on_first_page(diario_bytes, buffer_overlay)


def adicionar_assinatura_completa(diario_bytes, emissor, dados_assinatura, url_qrcode):
    """Adiciona assinatura visual no footer + QR code na primeira página (coluna esquerda)"""
    pdf_com_footer = adicionar_dados_assinatura_footer(diario_bytes, emissor, dados_assinatura)
    pdf_completo = adicionar_qrcode_assinatura_final(pdf_com_footer, url_qrcode, emissor, dados_assinatura)
    return pdf_completo


def remover_acentos(s: str) -> str:
    """
    Remove acentos de uma string.
    """
    return ''.join(
        c for c in unicodedata.normalize('NFD', s)
        if unicodedata.category(c) != 'Mn'
    )

import csv

def carregar_csv_em_dict():
    """Carrega nomes do CSV em um set (executado apenas uma vez no import do módulo)"""
    with open(os.path.join(DATA_DIR,'nomes.csv'), 'r', encoding='utf-8') as arquivo:
        leitor = csv.DictReader(arquivo)
        return {linha['nome'] for linha in leitor}

# Carregado uma única vez quando o módulo é importado
NOMES_CSV = carregar_csv_em_dict()    

def is_name(s: str) -> bool:
    """Verifica se a string começa com um nome válido do CSV"""
    if not s or not s.strip():
        return False
    
    # Remove acentos e pega apenas a primeira palavra
    s_sem_acentos = remover_acentos(s)
    primeira_palavra = s_sem_acentos.split(' ', 1)[0].capitalize()
    
    return primeira_palavra in NOMES_CSV

class DocumentReadError(Exception):
    def __init__(self, filename, original_exception):
        self.filename = filename
        self.original_exception = original_exception
